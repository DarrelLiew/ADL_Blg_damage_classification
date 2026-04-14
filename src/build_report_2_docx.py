"""
Build the DOCX report (4-experiment version, trial run omitted) for 60.001 ADL.

Usage:
    python build_report_2_docx.py

Outputs:
    ADL_Project_Report_2_v3.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
IMG = BASE / "images"
OUT = BASE / "ADL_Project_Report_2_v3.docx"


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D6E4F0')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption, width=6.5):
    path = IMG / filename
    if not path.exists():
        doc.add_paragraph(f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ══════════════════════════════════════════════════════════
    #  TITLE PAGE
    # ══════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("xView2 Building Damage Segmentation")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(25, 25, 25)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("A Siamese ResNet U-Net for Disaster Response")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("60.001 Applied Deep Learning, Y2026")
    run.bold = True
    run.font.size = Pt(12)

    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Singapore University of Technology and Design")
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("Team Members:", "[Member names and student IDs]"),
        ("Dataset:", "xView2 Tier 3 (Kaggle)"),
        ("GitHub:", "[Repository URL]"),
        ("Model Weights:", "[Google Drive / Dropbox link]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    #  ABSTRACT
    # ══════════════════════════════════════════════════════════
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "We built a Siamese ResNet U-Net that classifies building damage from satellite imagery "
        "into four levels: No Damage, Minor Damage, Major Damage, and Destroyed. Using the xView2 "
        "Tier 3 dataset (~2,200 image pairs), we achieved a damage macro F1 of 0.546 and mIoU of "
        "0.503. The main challenge is extreme class imbalance -- 97% of pixels are background. "
        "Through four experiments, we found that training strategy (weight decay, two-phase encoder "
        "freezing) mattered far more than architecture changes, accounting for the majority of our total "
        "improvement. Our Minor Damage F1 of 0.465 exceeds the typical range of xView2 competition "
        "winners (0.30-0.45), despite using only 10% of the full dataset with a single model."
    )

    # ══════════════════════════════════════════════════════════
    #  1. INTRODUCTION
    # ══════════════════════════════════════════════════════════
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "When disasters strike, rapid damage assessment saves lives. Satellite images become "
        "available within hours, but manually inspecting thousands of buildings does not scale. "
        "We built a deep learning model that takes a pair of satellite images (before and after "
        "a disaster) and outputs a pixel-level damage map."
    )
    doc.add_paragraph(
        "The main challenge is class imbalance: 97% of pixels are background. A model that "
        "predicts \"background\" everywhere gets 97% pixel accuracy but is useless. The hardest "
        "class is Minor Damage (only 0.3% of pixels) -- even competition winners only achieve "
        "0.30-0.45 F1 on it."
    )

    p = doc.add_paragraph()
    run = p.add_run("What we contribute:")
    run.bold = True

    for item in [
        "A Siamese ResNet U-Net with three-stream fusion that combines pre-disaster, post-disaster, and absolute difference features at every encoder level.",
        "A two-phase training strategy (freeze encoder first, then fine-tune with lower learning rate) that eliminated severe overfitting and gave us the biggest single improvement (+20.6% mIoU) without changing the architecture.",
        "Four systematic experiments where each one builds on the previous, showing that training strategy accounts for the majority of total improvement.",
        "A Minor Damage F1 of 0.465, beating the typical competition winner range despite having 10x less data.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "We improved from a baseline mIoU of 0.394 to 0.503 -- a 28% relative improvement "
        "across four experiments."
    )
    add_note(doc,
        "Note: We ran an initial 5-epoch trial run to verify the pipeline worked. Since it was "
        "too short to produce meaningful results, we omitted it from this report and start from "
        "our first full training run. Figure labels use the original experiment numbering "
        "(Exp 2-5 in figures = Exp 1-4 in this report)."
    )

    # ══════════════════════════════════════════════════════════
    #  2. DATASET
    # ══════════════════════════════════════════════════════════
    doc.add_heading("2. Dataset", level=1)

    doc.add_heading("2.1 Source", level=2)
    doc.add_paragraph(
        "The xView2 dataset [1] was created by the Defense Innovation Unit for the xView2 "
        "Building Damage Assessment Challenge. It contains WorldView-3 satellite imagery across "
        "multiple disaster types: earthquakes, hurricanes, wildfires, volcanic eruptions, and floods."
    )

    doc.add_heading("2.2 Structure", level=2)
    add_styled_table(doc, ["Property", "Value"], [
        ["Image pairs", "~2,200 (pre + post disaster)"],
        ["Original resolution", "1024 x 1024 px"],
        ["Training resolution", "512 x 512 px"],
        ["Annotation format", "GeoJSON polygons per building"],
        ["Split", "80% train / 10% val / 10% test (seed=42)"],
    ])

    doc.add_heading("2.3 Class Distribution", level=2)
    doc.add_paragraph("The dataset is extremely imbalanced:")
    add_styled_table(doc, ["Class", "Pixel %", "Description"], [
        ["Background", "~97%", "Non-building pixels"],
        ["No Damage", "~1.5%", "Intact buildings"],
        ["Minor Damage", "~0.3%", "Cracks, missing tiles"],
        ["Major Damage", "~0.5%", "Structural damage"],
        ["Destroyed", "~0.7%", "Collapsed buildings"],
    ])
    add_figure(doc, "C2_class_distribution.png",
        "Figure 1: Class distribution. Background dominates at ~97%. All four damage classes "
        "combined occupy less than 3% of pixels.")
    doc.add_paragraph(
        "Because of this imbalance, we use damage macro F1 (macro average of F1 for classes 1-4, "
        "excluding background) as our main metric instead of pixel accuracy."
    )

    doc.add_heading("2.4 Preprocessing", level=2)
    for item in [
        "Polygon annotations converted to pixel masks using Shapely and scikit-image",
        "Images resized from 1024x1024 to 512x512",
        "ImageNet normalisation applied (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("2.5 Data Augmentation", level=2)
    add_styled_table(doc, ["Augmentation", "Probability", "Purpose"], [
        ["Horizontal flip", "0.5", "Orientation invariance"],
        ["Vertical flip", "0.2", "Orientation invariance"],
        ["Random 90-degree rotation", "0.5", "Rotation invariance"],
        ["Brightness jitter", "+/-0.15", "Lighting robustness"],
        ["Contrast jitter", "+/-0.15", "Lighting robustness"],
        ["Saturation jitter", "+/-0.10", "Colour robustness"],
        ["Gaussian blur", "0.15", "Resolution robustness"],
        ["Gaussian noise", "0.15", "Sensor noise robustness"],
    ])

    # ══════════════════════════════════════════════════════════
    #  3. METHOD
    # ══════════════════════════════════════════════════════════
    doc.add_heading("3. Method", level=1)

    doc.add_heading("3.1 Architecture: Siamese ResNet U-Net", level=2)
    doc.add_paragraph(
        "Our model has three parts: a shared encoder, fusion blocks, and a decoder. Both the "
        "pre- and post-disaster images pass through the same ResNet encoder (pretrained on ImageNet). "
        "At each encoder level, a FusionBlock combines the two feature maps with an explicit change "
        "signal. A U-Net decoder then reconstructs the full-resolution 5-class damage segmentation map."
    )
    add_figure(doc, "C1_architecture_diagram.png",
        "Figure 2: Model architecture. Both images pass through the same ResNet encoder. "
        "FusionBlocks combine the features at each level. The U-Net decoder produces the "
        "final 5-class damage map.", width=6.2)

    p = doc.add_paragraph()
    run = p.add_run("Shared Encoder: ")
    run.bold = True
    p.add_run("Weight sharing ensures both images are in the same feature space. We tried two backbones:")
    add_styled_table(doc, ["Backbone", "Parameters", "Deepest Channels"], [
        ["ResNet34", "21M", "512"],
        ["ResNet50", "25M", "2048 (4x wider)"],
    ])

    p = doc.add_paragraph()
    run = p.add_run("Fusion Blocks: ")
    run.bold = True
    p.add_run(
        "At each encoder level: concat(pre, post, |post - pre|) -> ConvBlock. The absolute "
        "difference gives the model a direct signal of what changed between the two images."
    )

    p = doc.add_paragraph()
    run = p.add_run("U-Net Decoder: ")
    run.bold = True
    p.add_run(
        "Standard U-Net decoder with skip connections from fused features. Each level upsamples, "
        "concatenates the skip connection, and applies two Conv-BN-ReLU blocks. The final 1x1 "
        "convolution outputs 5 classes."
    )

    doc.add_heading("3.2 Loss Function", level=2)
    doc.add_paragraph("We use Focal Loss + Dice Loss (final configuration):")
    add_styled_table(doc, ["Component", "What it does"], [
        ["Focal Loss (gamma=2)", "Reduces influence of easy background pixels"],
        ["Dice Loss", "Measures region overlap, naturally handles class imbalance"],
        ["Log-inverse class weights", "Gives more importance to rare damage classes"],
        ["Label smoothing (0.05)", "Prevents overconfident predictions"],
    ])
    doc.add_paragraph(
        "Earlier experiments used Cross-Entropy instead of Focal Loss. We switched in Exp 4 "
        "because Major Damage F1 was dropping."
    )

    doc.add_heading("3.3 Two-Phase Training", level=2)
    doc.add_paragraph("This was our most important design decision:")
    p = doc.add_paragraph()
    run = p.add_run("Phase 1 (Frozen Encoder): ")
    run.bold = True
    p.add_run(
        "We freeze the ResNet encoder and only train the decoder and fusion blocks. This lets "
        "the decoder learn to work with the pretrained ImageNet features without messing them up."
    )
    p = doc.add_paragraph()
    run = p.add_run("Phase 2 (Fine-Tuning): ")
    run.bold = True
    p.add_run(
        "We unfreeze the top encoder layers with a much lower learning rate (6-30x lower than "
        "the decoder). This lets the encoder gradually adapt to satellite imagery."
    )
    add_figure(doc, "C3_two_phase_training.png",
        "Figure 3: Two-phase training. Phase 1 freezes the encoder. Phase 2 unfreezes top "
        "layers with a lower learning rate.", width=6.2)

    add_styled_table(doc, ["Parameter", "Exp 2", "Exp 3", "Exp 4"], [
        ["Freeze epochs", "15", "15", "10"],
        ["Unfrozen layers", "All", "2", "4"],
        ["Encoder LR", "1e-5", "1e-5", "5e-5"],
    ])

    doc.add_heading("3.4 Other Training Details", level=2)
    add_styled_table(doc, ["Parameter", "Value"], [
        ["Optimiser", "AdamW"],
        ["Decoder LR", "3e-4"],
        ["Weight decay", "0.05"],
        ["Gradient clipping", "1.0"],
        ["Batch size", "8 (effective 32-64 with grad. accum.)"],
        ["Mixed precision", "FP16"],
        ["Scheduler", "CosineAnnealingLR"],
        ["Early stopping", "Patience=15 on val damage F1"],
    ])

    # ══════════════════════════════════════════════════════════
    #  4. EXPERIMENTS
    # ══════════════════════════════════════════════════════════
    doc.add_heading("4. Experiments", level=1)
    doc.add_paragraph(
        "We ran four experiments. Each one was motivated by what we learned (or what went wrong) "
        "in the previous one."
    )
    add_figure(doc, "D1_prediction_grid.png",
        "Figure 4: Predictions across disaster types. The model gets progressively better at "
        "identifying damage across experiments.", width=6.5)

    # ── Exp 1: Baseline ──
    doc.add_heading("4.1 Experiment 1: Baseline", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: ")
    run.bold = True
    p.add_run("Establish baseline performance with a full training run.")

    doc.add_paragraph(
        "Setup: ResNet34 backbone, all layers unfrozen, effective batch size 2, 50 epochs "
        "(ran 35 before stopping)."
    )
    doc.add_paragraph(
        "Results: mIoU = 0.394, but with severe overfitting -- train loss dropped to 0.63 while "
        "val loss climbed to 2.38 (a 3.8x gap)."
    )
    add_figure(doc, "A1_exp2_overfitting.png",
        "Figure 5: Exp 1 train vs val loss. The gap between the curves shows severe overfitting.",
        width=5.5)

    p = doc.add_paragraph()
    run = p.add_run("Why the val loss is noisy. ")
    run.bold = True
    p.add_run(
        "The cosine LR scheduler was set to restart every few epochs rather than run one long "
        "decay to the end of training. Each LR restart momentarily shoves the weights out of "
        "their current local minimum, which spikes validation loss before the next cosine phase "
        "re-settles. Combined with weak weight decay (1e-4) -- which let the model memorise "
        "noise -- every restart reveals a different overfitted state. The train loss stays smooth "
        "because it is measured on the memorised set; val loss is not, so it fluctuates visibly."
    )

    p = doc.add_paragraph()
    run = p.add_run("Takeaway: ")
    run.bold = True
    p.add_run("The model was memorising the training data. We needed a fundamentally different approach.")

    # ── Exp 2: Training Strategy Overhaul ──
    doc.add_heading("4.2 Experiment 2: Training Configuration Changes", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: ")
    run.bold = True
    p.add_run("Reduce overfitting by modifying training hyperparameters and strategy, without changing the architecture.")

    add_styled_table(doc, ["Change", "Before", "After", "Why"], [
        ["Weight decay", "1e-4", "0.05 (500x more)", "Main overfitting fix"],
        ["Encoder", "All unfrozen", "Phase 1 frozen, Phase 2 unfrozen", "Protect pretrained features"],
        ["Encoder LR", "Same as decoder", "30x lower", "Preserve ImageNet knowledge"],
        ["Effective batch", "2", "8", "More stable gradients"],
        ["Class weights", "Basic freq. inverse", "Log-inverse (all samples)", "Better for rare classes"],
        ["Metric", "mIoU", "Damage macro F1", "Ignores background"],
    ])

    p = doc.add_paragraph()
    run = p.add_run("Results: ")
    run.bold = True
    p.add_run("mIoU = 0.475 (+20.6%), Damage F1 = 0.510")

    doc.add_paragraph(
        "This was our biggest improvement -- and we changed nothing about the architecture. The "
        "500x weight decay increase was the main driver: the train-val loss gap closed from "
        "3.8x (train 0.63 vs val 2.38) down to ~1:1 (train 2.48 vs val 2.47) by the end of "
        "training. Train and val loss now move together; the model has stopped memorising."
    )
    p = doc.add_paragraph()
    run = p.add_run("Why the absolute loss values went up. ")
    run.bold = True
    p.add_run(
        "Exp 2's train/val losses (2.48 each) look worse than Exp 1's train loss (0.63), but "
        "that comparison is misleading. Exp 2 uses log-inverse class weights and label smoothing, "
        "which inflate the raw loss by penalising rare-class mistakes more heavily. The "
        "segmentation quality is much better (mIoU 0.394 -> 0.475), which is what matters."
    )
    add_figure(doc, "A2_exp2v3_overfitting_fix.png",
        "Figure 6: Before vs after. Left: Exp 1 with diverging train/val loss. Right: Exp 2 "
        "with parallel curves. Zero architecture changes.", width=6.2)

    doc.add_paragraph("Per-class F1:")
    add_styled_table(doc, ["Class", "F1"], [
        ["No Damage", "0.672"], ["Minor Damage", "0.283"],
        ["Major Damage", "0.461"], ["Destroyed", "0.623"],
    ])

    p = doc.add_paragraph()
    run = p.add_run("New problem: ")
    run.bold = True
    p.add_run(
        "Minor Damage F1 was only 0.283 -- much worse than other classes. Maybe ResNet34 doesn't "
        "have enough capacity to detect subtle damage like cracks and missing tiles?"
    )

    # ── Exp 3: ResNet50 ──
    doc.add_heading("4.3 Experiment 3: ResNet50 Backbone", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: ")
    run.bold = True
    p.add_run("Use a bigger encoder to improve Minor Damage detection.")

    doc.add_paragraph("Change: Swapped ResNet34 (21M params) for ResNet50 (25M params, 4x wider features).")

    p = doc.add_paragraph()
    run = p.add_run("Results: ")
    run.bold = True
    p.add_run("mIoU = 0.488 (+2.7%), Damage F1 = 0.521 (+2.2%)")

    doc.add_paragraph("The overall numbers improved slightly, but the per-class results told a different story:")
    add_styled_table(doc, ["Class", "Exp 2", "Exp 3", "Change"], [
        ["No Damage", "0.672", "0.709", "+0.037"],
        ["Minor Damage", "0.283", "0.433", "+0.150 (+53%)"],
        ["Major Damage", "0.461", "0.384", "-0.077 (-17%)"],
        ["Destroyed", "0.623", "0.557", "-0.066"],
    ])
    add_figure(doc, "A3_exp3v4_perclass.png",
        "Figure 7: Per-class F1 comparison. Minor Damage jumped +53% with ResNet50, but Major "
        "Damage dropped -17%.", width=5.5)

    doc.add_paragraph(
        "Minor Damage improved because ResNet50 has wider features that can capture subtle cracks. "
        "But Major Damage and Destroyed got worse -- with encoder_lr=1e-5 and only 2 layers "
        "unfrozen, the ResNet50 backbone wasn't adapting enough to satellite imagery."
    )
    p = doc.add_paragraph()
    run = p.add_run("Lesson: ")
    run.bold = True
    p.add_run("A bigger model isn't automatically better. You have to let it learn.")

    # ── Exp 4: Tuned Fine-Tuning + Focal ──
    doc.add_heading("4.4 Experiment 4: Tuned Fine-Tuning + Focal Loss", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: ")
    run.bold = True
    p.add_run("Let ResNet50 actually adapt to satellite imagery.")

    add_styled_table(doc, ["Change", "Before", "After", "Why"], [
        ["Encoder LR", "1e-5", "5e-5 (5x higher)", "Let encoder adapt faster"],
        ["Freeze epochs", "15", "10", "Decoder converges earlier"],
        ["Unfrozen layers", "2", "4", "Deeper adaptation"],
        ["Loss", "CE + Dice", "Focal (gamma=2) + Dice", "Focus on hard boundary pixels"],
    ])

    doc.add_paragraph("Results (full 60-epoch run; best epoch 49):")
    add_styled_table(doc, ["Metric", "Value", "Change vs Exp 3"], [
        ["mIoU", "0.503", "+0.015"],
        ["Damage F1", "0.546", "+0.025"],
    ])

    add_styled_table(doc, ["Class", "Exp 3", "Exp 4", "Change"], [
        ["No Damage", "0.709", "0.692", "-0.017"],
        ["Minor Damage", "0.433", "0.465", "+0.032"],
        ["Major Damage", "0.384", "0.437", "+0.053"],
        ["Destroyed", "0.557", "0.588", "+0.031"],
    ])

    doc.add_paragraph(
        "The higher encoder LR and more unfrozen layers let ResNet50 actually learn "
        "satellite-specific features. Focal Loss helped by focusing on hard boundary pixels, "
        "improving every damage class over Exp 3."
    )
    add_figure(doc, "A4_exp4v5_phase2_adaptation.png",
        "Figure 8: Exp 3 vs Exp 4 damage F1 during Phase 2. More aggressive fine-tuning "
        "pushed Exp 4 higher.", width=5.8)
    add_figure(doc, "A5_exp5_deepdive.png",
        "Figure 9: Exp 4 deep dive. (a) Damage F1 over all 60 epochs with phase boundary. "
        "(b) Per-class F1 trajectories.", width=6.0)

    p = doc.add_paragraph()
    run = p.add_run("Why the validation curves are so noisy. ")
    run.bold = True
    p.add_run(
        "Phase 2 validation F1 oscillates by 0.10-0.20 between adjacent epochs, with several "
        "sharp drops (e.g. epochs 19-20, 22-23, 25-26, 37-38, 40, 54-55, 57). Two things cause this:"
    )
    for item in [
        "Aggressive encoder fine-tuning. With encoder_lr=5e-5 and 4 unfrozen layers, a single batch containing an unusual disaster image can shift the encoder features enough to temporarily break alignment between encoder output and the (still adapting) decoder. The next epoch usually recovers once the optimizer re-balances.",
        "Focal Loss + log-inverse class weights concentrate gradient on rare pixels. The damage classes each occupy under 1% of pixels, so a handful of wrongly-classified buildings in one validation batch can swing the per-class F1 and move the macro average visibly. Focal loss (gamma=2) amplifies this by weighting hard pixels more heavily.",
    ]:
        doc.add_paragraph(item, style='List Number')
    doc.add_paragraph(
        "This is not overfitting. Training loss stays flat around 2.79-2.81 through Phase 2 "
        "while val loss spikes, and the best val F1 keeps climbing -- the trend is noisy but "
        "upward, which is why the best epoch (49) comes late in the schedule."
    )

    # ── Results Summary ──
    doc.add_heading("4.5 Results Summary", level=2)
    add_styled_table(doc,
        ["", "Exp 1", "Exp 2", "Exp 3", "Exp 4"],
        [
            ["Key Change", "Baseline", "Training overhaul", "ResNet50", "Tuned fine-tuning + Focal"],
            ["Backbone", "ResNet34", "ResNet34", "ResNet50", "ResNet50"],
            ["Epochs", "35", "42", "32", "60"],
            ["Val mIoU", "0.394", "0.475", "0.488", "0.503"],
            ["Val Damage F1", "--", "0.510", "0.521", "0.546"],
            ["Overfitting", "Severe", "Controlled", "Controlled", "Controlled"],
        ])

    add_figure(doc, "B5_combined_progression.png",
        "Figure 10: mIoU and Damage F1 progression across all experiments.", width=6.0)
    add_figure(doc, "D2_improvement_closeup.png",
        "Figure 11: Close-up comparison of predictions. Damage classification improves across experiments.", width=6.2)
    add_figure(doc, "C4_loss_landscape_all.png",
        "Figure 12: Train and val loss across all experiments. Overfitting decreases progressively.", width=6.0)

    # ── SOTA ──
    doc.add_heading("4.6 Comparison to State-of-the-Art", level=2)
    add_styled_table(doc, ["Model", "Damage F1", "Data", "Ensemble"], [
        ["1st place (vdurnov) [5]", "~0.75", "Full xBD (22K pairs)", "12 models"],
        ["2nd place (selimsef) [6]", "~0.72", "Full xBD", "2 models"],
        ["Single-model baseline", "~0.625", "Full xBD", "1 model"],
        ["Ours (Exp 4)", "0.546", "Tier 3 (2.2K pairs)", "1 model"],
    ])
    add_figure(doc, "B4_sota_comparison.png",
        "Figure 13: Comparison to competition results. The gap comes from having less data, no "
        "ensembling, and no localisation pretraining.", width=5.8)

    doc.add_paragraph("The gap is due to resource differences, not architecture:")
    for item in [
        "10x less data -- we used Tier 3 (2.2K pairs) vs the full xBD (22K pairs)",
        "No localisation pretraining -- winners first trained a model to find buildings, then fine-tuned for damage",
        "Single model -- no ensembling or test-time augmentation",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        "Despite this, our 28% relative mIoU improvement shows effective hyperparameter tuning, "
        "and our Minor Damage F1 (0.465) beats the typical competition winner range (0.30-0.45)."
    )

    # ══════════════════════════════════════════════════════════
    #  5. ANALYSIS
    # ══════════════════════════════════════════════════════════
    doc.add_heading("5. Analysis", level=1)

    doc.add_heading("5.1 What Helped Most", level=2)
    add_figure(doc, "B1_miou_waterfall.png",
        "Figure 14: Waterfall chart of mIoU improvement. Training strategy (Exp 2) accounts "
        "for the majority of the total gain.", width=6.0)
    add_styled_table(doc, ["Rank", "Change", "Experiment", "mIoU Gain", "% of Total"], [
        ["1", "Training strategy overhaul", "Exp 2", "+0.081", "74%"],
        ["2", "Focal Loss + tuned fine-tuning", "Exp 4", "+0.015", "14%"],
        ["3", "Bigger backbone (ResNet50)", "Exp 3", "+0.013", "12%"],
    ])
    p = doc.add_paragraph()
    run = p.add_run("Key takeaway: ")
    run.bold = True
    p.add_run("How you train matters more than what you train. The 500x weight decay increase was the single most impactful change.")

    doc.add_heading("5.2 Minor Damage: The Hardest Class", level=2)
    add_figure(doc, "B2_rare_class_challenge.png",
        "Figure 15: Per-class F1 across Exp 2-4. Minor Damage improved by +65% over three "
        "experiments and now exceeds the competition winner range.", width=5.8)
    add_styled_table(doc, ["Experiment", "Minor Damage F1", "Change"], [
        ["Exp 2 (ResNet34)", "0.283", "--"],
        ["Exp 3 (ResNet50)", "0.433", "+53%"],
        ["Exp 4 (+ Focal)", "0.465", "+64% total"],
    ])
    doc.add_paragraph(
        "Minor Damage is hard because cracks and missing tiles are tiny at 512x512 resolution. "
        "Competition winners also struggle with this class (typically 0.30-0.45 F1). Our final "
        "0.465 exceeds this range."
    )

    doc.add_heading("5.3 Bigger Backbone Needs Bigger Learning Rate", level=2)
    add_figure(doc, "B3_encoder_adaptation_arc.png",
        "Figure 16: Per-class F1 across experiments. Major Damage dropped when we added ResNet50 "
        "with conservative fine-tuning, then recovered when we increased the encoder learning rate.",
        width=5.5)
    doc.add_paragraph(
        "Switching to ResNet50 (Exp 3) actually made Major Damage worse (0.461 -> 0.384). This "
        "happened because the encoder learning rate was too low. When we increased the encoder LR "
        "5x and unfroze more layers in Exp 4, Major Damage recovered to 0.437."
    )
    p = doc.add_paragraph()
    run = p.add_run("Lesson: A bigger model with weak fine-tuning can be worse than a smaller model that is fully adapted.")
    run.bold = True

    doc.add_heading("5.4 What Didn't Help", level=2)
    for item in [
        "Baseline without regularisation (Exp 1): 35 epochs with default settings resulted in severe overfitting.",
        "Conservative fine-tuning (Exp 3): Keeping encoder LR at 1e-5 with only 2 layers unfrozen wasted ResNet50's extra capacity.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("5.5 Failure Cases", level=2)
    add_figure(doc, "D3_failure_cases.png",
        "Figure 17: Where Exp 4 still fails. Error maps show misclassifications for minor/major "
        "confusion, small buildings, and mixed damage levels.", width=6.0)
    for item in [
        "Minor vs Major confusion: The boundary between damage levels is subjective -- even human annotators disagree.",
        "Small buildings: Buildings that are only a few pixels wide are hard to classify at 512x512 resolution.",
        "Underrepresented disasters: The model performs worse on disaster types with fewer training examples.",
        "Phase 2 instability: Aggressive fine-tuning caused occasional validation loss spikes throughout the 60-epoch run, though the model always recovered. These likely contributed to the late-epoch plateau after the best epoch (49).",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ══════════════════════════════════════════════════════════
    #  6. REPRODUCIBILITY
    # ══════════════════════════════════════════════════════════
    doc.add_heading("6. Reproducibility", level=1)

    doc.add_heading("6.1 Dependencies", level=2)
    doc.add_paragraph(
        "Python 3.11+, PyTorch >= 2.0 (with CUDA), torchvision, opencv-python, scikit-image, "
        "scikit-learn, shapely, matplotlib, numpy, tqdm."
    )

    doc.add_heading("6.2 Dataset Setup", level=2)
    for item in [
        "Download xView2 Tier 3 from Kaggle",
        "Place in: data/train/train/images/ and data/train/train/labels/",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("6.3 Training", level=2)
    for item in [
        "Open Kaggle-latest.ipynb (Kaggle) or training_script.ipynb (local)",
        "Set CONFIG_PATH to the experiment config (e.g. training_runs/siamese_resunet_xview2_4/config.json)",
        "Run all cells. Outputs saved automatically: config.json, history.json, metrics.csv, best_model.pt, last_model.pt",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("6.4 Loading a Trained Model", level=2)
    doc.add_paragraph(
        "model = SiameseResUNet(backbone='resnet50', num_classes=5, decoder_channels=[512, 256, 128, 64])\n"
        "checkpoint = torch.load('training_runs/siamese_resunet_xview2_4/best_model.pt', map_location='cpu')\n"
        "model.load_state_dict(checkpoint['model_state_dict'])\n"
        "model.eval()\n"
        "pred = model(pre_image, post_image).argmax(dim=1)"
    )

    doc.add_heading("6.5 Recreating Figures", level=2)
    doc.add_paragraph("python generate_figures.py")
    doc.add_paragraph("python generate_extra_figures.py")
    doc.add_paragraph("python generate_prediction_figures.py")

    doc.add_heading("6.6 Project Structure", level=2)
    doc.add_paragraph(
        "ADL Project/\n"
        "  REPORT_2.md                  # This report (markdown version)\n"
        "  Kaggle-latest.ipynb          # Training notebook\n"
        "  training_script.ipynb        # Local training notebook\n"
        "  generate_figures.py          # Chart figures (A1-B5, C2, C4)\n"
        "  generate_extra_figures.py    # Diagram figures (C1, C3)\n"
        "  generate_prediction_figures.py  # Prediction figures (D1-D3)\n"
        "  images/                      # 17 figures\n"
        "  training_runs/\n"
        "    siamese_resunet_xview2_1/  # Initial trial (omitted from report)\n"
        "    siamese_resunet_xview2_2/  # Exp 1: Baseline\n"
        "    siamese_resunet_xview2/    # Exp 2: Training overhaul\n"
        "    siamese_resunet_xview2_3/  # Exp 3: ResNet50\n"
        "    siamese_resunet_xview2_4/  # Exp 4: Focal + Tuned Phase 2\n"
        "  data/                        # Dataset (not in git)"
    )

    # ══════════════════════════════════════════════════════════
    #  7. CONCLUSION
    # ══════════════════════════════════════════════════════════
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "We built a Siamese ResNet U-Net for satellite building damage segmentation, reaching a "
        "damage macro F1 of 0.546 and mIoU of 0.503 across four experiments. Three main findings:"
    )
    findings = [
        ("1. Training strategy matters more than architecture. ",
         "The Exp 2 training overhaul gave us +20.6% mIoU with zero architecture changes. "
         "Weight decay (0.05) and two-phase encoder freezing accounted for the majority of "
         "our total improvement."),
        ("2. Bigger models need proper fine-tuning. ",
         "ResNet50 improved Minor Damage by +53%, but hurt Major Damage when fine-tuned too "
         "conservatively. Increasing the encoder learning rate 5x and unfreezing more layers fixed this."),
        ("3. Focal Loss helps hard cases. ",
         "Switching to Focal Loss directed the model toward difficult damage boundary pixels, "
         "recovering the Major Damage regression and pushing Minor Damage to 0.465 -- above "
         "the competition winner range."),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(body)

    doc.add_paragraph(
        "The gap between our result (0.546) and competition winners (0.75) comes from resource "
        "constraints: 10x less data, no ensembling, and no localisation pretraining."
    )

    doc.add_heading("Future Work", level=2)
    for item in [
        "Add test-time augmentation (typically +1-3% for free)",
        "Add building localisation head (multi-task learning)",
        "Train on the full xBD dataset if compute allows",
        "Experiment with model ensembling (different seeds or architectures)",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ══════════════════════════════════════════════════════════
    #  8. TEAM CONTRIBUTIONS
    # ══════════════════════════════════════════════════════════
    doc.add_heading("8. Team Contributions", level=1)
    add_styled_table(doc, ["Member", "Contributions"], [
        ["[Name 1]", "[Contributions]"],
        ["[Name 2]", "[Contributions]"],
        ["[Name 3]", "[Contributions]"],
        ["[Name 4]", "[Contributions]"],
    ])

    # ══════════════════════════════════════════════════════════
    #  REFERENCES
    # ══════════════════════════════════════════════════════════
    doc.add_heading("References", level=1)
    refs = [
        '[1] Gupta, R., et al. "xBD: A Dataset for Assessing Building Damage from Satellite Imagery." arXiv:1911.09296, 2019.',
        '[2] Ronneberger, O., Fischer, P., & Brox, T. "U-Net: Convolutional Networks for Biomedical Image Segmentation." MICCAI 2015.',
        '[3] He, K., et al. "Deep Residual Learning for Image Recognition." CVPR 2016.',
        '[4] Lin, T.-Y., et al. "Focal Loss for Dense Object Detection." ICCV 2017.',
        '[5] xView2 1st Place Solution: github.com/vdurnov/xview2_1st_place_solution',
        '[6] xView2 2nd Place Solution: github.com/selimsef/xview2_solution',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)

    # ── Save ──
    doc.save(str(OUT))
    print(f"Report saved to: {OUT}")


if __name__ == "__main__":
    build()
