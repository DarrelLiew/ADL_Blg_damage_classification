"""
Build the final DOCX report for 60.001 Applied Deep Learning.

Usage:
    python build_report_docx.py

Outputs:
    ADL_Project_Report.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).parent
IMG = BASE / "images"
OUT = BASE / "ADL_Project_Report.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D6E4F0')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing
    return table


def add_figure(doc, filename, caption, width=6.5):
    """Add an image with caption."""
    path = IMG / filename
    if not path.exists():
        doc.add_paragraph(f"[Figure not found: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = doc.styles['Caption'] if 'Caption' in [s.name for s in doc.styles] else None
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()  # spacing


def build():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ╔════════════════════════════════════════════════════════════╗
    # ║  TITLE PAGE                                                ║
    # ╚════════════════════════════════════════════════════════════╝
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("xView2 Building Damage Segmentation")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(25, 25, 25)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI for Disaster Response")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("60.001 Applied Deep Learning, Y2026")
    run.bold = True
    run.font.size = Pt(12)

    school = doc.add_paragraph()
    school.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school.add_run("Singapore University of Technology and Design")
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph()

    details = [
        ("Team Members:", "[Member names and student IDs]"),
        ("Dataset:", "xView2 Tier 3 (Kaggle)"),
        ("GitHub:", "[Repository URL]"),
        ("Model Weights:", "[Google Drive / Dropbox link]"),
    ]
    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(value)
        run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Submitted: April 17, 2026")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.add_page_break()

    # ╔════════════════════════════════════════════════════════════╗
    # ║  1. INTRODUCTION                                           ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1 Problem Statement", level=2)
    doc.add_paragraph(
        "When natural disasters strike, rapid assessment of building damage is critical for directing "
        "emergency resources. Traditional damage assessment relies on ground surveys that take days to "
        "weeks -- time that costs lives. Satellite imagery is available within hours, but manually "
        "inspecting thousands of buildings across disaster zones does not scale."
    )
    doc.add_paragraph(
        "This project develops a deep learning model for automated building damage segmentation from "
        "pre- and post-disaster satellite imagery. Given a pair of satellite images (before and after a "
        "disaster), the model produces a pixel-level damage map classifying every building into one of "
        "four damage levels: No Damage, Minor Damage, Major Damage, and Destroyed."
    )

    doc.add_heading("1.2 Objectives", level=2)
    for obj in [
        "Design and implement a Siamese U-Net architecture for change-detection-based damage segmentation",
        "Systematically improve model performance through iterative experimentation -- training strategy, regularisation, backbone scaling, and loss function design",
        "Evaluate against the xView2 competition benchmark and analyse failure modes",
        "Deliver a reproducible, well-documented codebase with saved model weights",
    ]:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading("1.3 Scope and Constraints", level=2)
    for item in [
        "Dataset: xView2 Tier 3 (~2,200 image pairs), a subset of the full xBD dataset (~22,000 pairs)",
        "Compute: Single GPU training (Kaggle P100 / SUTD AI Cluster). GPU quota limitations and cluster downtime affected our final experiment.",
        "Architecture: Single-model inference (no ensembling), practical for real deployment",
        "Classes: 5-class segmentation -- Background (0), No Damage (1), Minor Damage (2), Major Damage (3), Destroyed (4)",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("1.4 Social Impact", level=2)
    doc.add_paragraph(
        "Automated damage assessment directly supports humanitarian response. It reduces assessment time "
        "from days to minutes, can process entire cities simultaneously, and provides consistent "
        "classification across analysts. Organisations like the UN, Red Cross, and FEMA have identified "
        "automated satellite damage assessment as a priority capability for disaster response."
    )

    # ╔════════════════════════════════════════════════════════════╗
    # ║  2. DATASET                                                ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("2. Dataset", level=1)

    doc.add_heading("2.1 Source", level=2)
    doc.add_paragraph(
        "The xView2 dataset was created by the Defense Innovation Unit (DIU) for the xView2 Building "
        "Damage Assessment Challenge. It contains WorldView-3 satellite imagery (1.4 m GSD) across "
        "multiple disaster types: earthquakes, hurricanes, wildfires, volcanic eruptions, and floods."
    )

    doc.add_heading("2.2 Structure", level=2)
    add_styled_table(doc,
        ["Property", "Value"],
        [
            ["Image pairs", "~2,200 (pre + post disaster)"],
            ["Original resolution", "1024 x 1024 px"],
            ["Training resolution", "512 x 512 px (resized)"],
            ["Annotation format", "GeoJSON polygons per building"],
            ["Split", "80% train / 10% val / 10% test"],
            ["Random seed", "42 (deterministic split)"],
        ])

    doc.add_heading("2.3 Class Distribution", level=2)
    doc.add_paragraph(
        "The dataset exhibits extreme class imbalance, which is the central challenge of this task. "
        "Background pixels constitute approximately 97% of all pixels. The four damage classes share "
        "the remaining ~3%, with Minor Damage being the rarest at only 0.3% of pixels."
    )
    add_figure(doc, "C2_class_distribution.png",
        "Figure 1: xView2 Tier 3 class distribution. Left: full scale showing background dominance. "
        "Right: zoomed view of damage classes. Minor Damage (0.3%) is the hardest class to detect.")

    doc.add_paragraph(
        "This imbalance means a model predicting 'background' everywhere achieves 97% pixel accuracy "
        "while being completely useless for damage assessment. This motivated our adoption of Damage "
        "Macro F1 (macro average of F1 for damage classes 1-4, excluding background) as the primary "
        "evaluation metric from Experiment 3 onward."
    )

    doc.add_heading("2.4 Preprocessing and Augmentation", level=2)
    doc.add_paragraph(
        "Polygon annotations were rasterised to per-pixel masks using Shapely and scikit-image. Images "
        "were resized from 1024x1024 to 512x512 via bilinear interpolation. ImageNet normalisation was "
        "applied (mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])."
    )
    doc.add_paragraph("Data augmentation (final configuration, Exp 5):")
    add_styled_table(doc,
        ["Augmentation", "Probability", "Purpose"],
        [
            ["Horizontal flip", "0.5", "Orientation invariance"],
            ["Vertical flip", "0.2", "Orientation invariance"],
            ["Random 90-degree rotation", "0.5", "Rotation invariance"],
            ["Brightness/Contrast jitter", "+/-0.15", "Lighting robustness"],
            ["Saturation jitter", "+/-0.10", "Colour robustness"],
            ["Gaussian blur", "0.15", "Resolution robustness"],
            ["Gaussian noise", "0.15", "Sensor noise robustness"],
        ])

    # ╔════════════════════════════════════════════════════════════╗
    # ║  3. METHOD                                                 ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("3. Method", level=1)

    doc.add_heading("3.1 Architecture: Siamese ResNet U-Net", level=2)
    doc.add_paragraph(
        "The model follows a Siamese encoder -- fusion -- decoder paradigm designed for change "
        "detection in satellite imagery. A shared ResNet backbone extracts hierarchical features from "
        "both the pre- and post-disaster images. At each encoder level, a FusionBlock combines the two "
        "feature maps with an explicit change signal. A U-Net decoder then reconstructs the full-resolution "
        "5-class damage segmentation map."
    )
    add_figure(doc, "C1_architecture_diagram.png",
        "Figure 2: Siamese ResNet U-Net architecture. Pre- and post-disaster images are processed by a "
        "shared ResNet encoder. FusionBlocks at each level combine features via 3-stream fusion "
        "(pre + post + |diff|). A U-Net decoder with skip connections produces the 5-class damage map.",
        width=6.2)

    p = doc.add_paragraph()
    run = p.add_run("Shared Encoder: ")
    run.bold = True
    p.add_run(
        "A pretrained ResNet backbone (ImageNet weights) extracts hierarchical features from both images. "
        "Weight sharing ensures the same feature space for both temporal views."
    )
    add_styled_table(doc,
        ["Backbone", "Parameters", "Deepest Channels", "Block Type"],
        [["ResNet34", "21M", "512", "BasicBlock"], ["ResNet50", "25M", "2048", "Bottleneck (4x wider)"]])

    p = doc.add_paragraph()
    run = p.add_run("Fusion Blocks: ")
    run.bold = True
    p.add_run(
        "At each encoder level, a FusionBlock combines pre- and post-disaster features using three-stream "
        "fusion: concat(pre_feat, post_feat, |post_feat - pre_feat|). The absolute difference explicitly "
        "encodes change magnitude. The concatenated 3C-channel tensor is compressed to C channels by a ConvBlock.")

    p = doc.add_paragraph()
    run = p.add_run("U-Net Decoder: ")
    run.bold = True
    p.add_run(
        "Standard U-Net decoder with skip connections from fused features. Each DecoderBlock: "
        "bilinear upsample -> concatenate skip -> ConvBlock. A 1x1 convolution produces the 5-class output.")

    doc.add_heading("3.2 Loss Function", level=2)
    doc.add_paragraph("Our loss function evolved across experiments. The final configuration (Exp 5) uses:")
    p = doc.add_paragraph()
    run = p.add_run("L = Focal Loss(pred, target, class_weights, gamma=2) + Dice Loss(pred, target)")
    run.bold = True

    add_styled_table(doc,
        ["Component", "Purpose"],
        [
            ["Focal Loss (gamma=2)", "Downweights easy background pixels; focuses on hard boundaries"],
            ["Dice Loss", "Region-level overlap; inherently handles class imbalance"],
            ["Log-inverse class weights", "Upweights rare damage classes (clipped to [1, 20])"],
            ["Label smoothing (0.05)", "Prevents overconfident predictions on easy pixels"],
        ])

    doc.add_heading("3.3 Two-Phase Training Strategy", level=2)
    doc.add_paragraph(
        "A key contribution of this project is the two-phase training strategy, developed in Exp 3 "
        "after observing severe overfitting in Exp 2.")
    add_figure(doc, "C3_two_phase_training.png",
        "Figure 3: Two-phase training strategy. Phase 1 freezes the encoder so the decoder and fusion "
        "blocks learn to interpret pretrained features. Phase 2 unfreezes encoder layers with a lower "
        "learning rate for domain-specific adaptation.", width=6.2)

    doc.add_paragraph(
        "Phase 1 (Frozen Encoder): Encoder weights frozen; only decoder and fusion blocks train at lr=3e-4. "
        "Duration: 10-15 epochs.")
    doc.add_paragraph(
        "Phase 2 (Full Fine-Tuning): Top N encoder layers unfrozen with differential LR 6-30x lower "
        "than decoder. Allows task-specific adaptation to satellite imagery.")

    add_styled_table(doc,
        ["Parameter", "Exp 3", "Exp 4", "Exp 5"],
        [["Freeze epochs", "15", "15", "10"], ["Unfrozen layers", "All", "2", "4"], ["Encoder LR", "1e-5", "1e-5", "5e-5"]])

    doc.add_heading("3.4 Optimisation Details", level=2)
    add_styled_table(doc,
        ["Parameter", "Value"],
        [
            ["Optimiser", "AdamW"],
            ["Decoder LR", "3e-4"],
            ["Encoder LR (Phase 2)", "1e-5 to 5e-5 (tuned per experiment)"],
            ["Weight decay", "0.05"],
            ["Gradient clipping", "1.0 (max norm)"],
            ["Batch size", "8-16 (final config)"],
            ["Gradient accumulation", "4-8 steps (effective batch 32-64)"],
            ["Mixed precision", "FP16 (torch.cuda.amp)"],
            ["Scheduler", "CosineAnnealingLR (min_lr=1e-6)"],
            ["Early stopping", "Patience=15 on val damage F1"],
        ])

    # ╔════════════════════════════════════════════════════════════╗
    # ║  4. EXPERIMENTS                                            ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("4. Experiments", level=1)
    doc.add_paragraph(
        "We conducted 5 experiments, each building on the previous. This iterative approach allowed us "
        "to isolate the impact of individual changes and understand what drives performance. Each "
        "experiment was motivated by a specific observation or failure from the previous one.")
    doc.add_paragraph(
        "The figure below shows model predictions across four disaster types for three key experiments. "
        "The visual progression from near-blank predictions in Exp 1 to structured damage maps in Exp 5 "
        "mirrors the quantitative improvements reported below.")
    add_figure(doc, "D1_prediction_grid.png",
        "Figure 15: Prediction comparison across 4 disaster types and 3 experiments. "
        "Exp 1 produces near-blank or noisy predictions. Exp 3 begins detecting damage. "
        "Exp 5 produces structured damage maps approaching ground truth quality.", width=6.5)

    # ── Exp 1 ──
    doc.add_heading("4.1 Experiment 1: Baseline", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: "); run.bold = True
    p.add_run("Get a working model and establish baseline performance.")

    add_styled_table(doc, ["Parameter", "Value"],
        [["Backbone", "ResNet34 (all layers unfrozen)"], ["Batch size", "1 (effective 2)"],
         ["Epochs", "5"], ["LR", "1e-4, CosineAnnealingLR"], ["Weight decay", "1e-4"],
         ["Loss", "CE + Dice (equal weight)"], ["Class weights", "Frequency inverse (256 samples)"]])

    add_styled_table(doc, ["Metric", "Best Value", "Epoch"],
        [["Val mIoU", "0.359", "2"], ["Val Mean Dice", "0.443", "2"], ["Val Pixel Accuracy", "0.944", "2"]])

    p = doc.add_paragraph()
    run = p.add_run("Outcome: "); run.bold = True
    p.add_run(
        "94.4% pixel accuracy was misleading -- the model was good at predicting background. "
        "Overfitting appeared by epoch 3. Cosine LR decayed to near-zero within 5 epochs.")
    p = doc.add_paragraph()
    run = p.add_run("Takeaway: "); run.bold = True
    p.add_run("Baseline mIoU=0.359. Was the problem insufficient epochs, or something deeper?")

    # ── Exp 2 ──
    doc.add_heading("4.2 Experiment 2: More Epochs", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: "); run.bold = True
    p.add_run("Determine whether the baseline simply needed more training time.")
    doc.add_paragraph("Change: Epochs 5 -> 50 (ran 35). All else identical to Exp 1.")

    add_styled_table(doc, ["Metric", "Best Value", "Epoch", "vs. Exp 1"],
        [["Val mIoU", "0.394", "19", "+0.035 (+9.7%)"], ["Val Mean Dice", "0.481", "35", "+0.038 (+8.6%)"]])

    p = doc.add_paragraph()
    run = p.add_run("Outcome: "); run.bold = True
    p.add_run(
        "Severe overfitting dominated. By epoch 35, train loss 0.63 vs val loss 2.38 (3.8:1 ratio). "
        "Val mIoU peaked at epoch 19 then fluctuated for 16 more epochs.")
    add_figure(doc, "A1_exp2_overfitting.png",
        "Figure 4: Exp 2 train vs val loss showing severe overfitting (3.8:1 gap).", width=5.5)

    p = doc.add_paragraph()
    run = p.add_run("Takeaway: "); run.bold = True
    p.add_run("More time was not the answer. The model needed regularisation and a new training approach.")

    # ── Exp 3 ──
    doc.add_heading("4.3 Experiment 3: Training Strategy Overhaul -- The Breakthrough", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: "); run.bold = True
    p.add_run("Eliminate overfitting and force the model to learn damage classes. Zero architecture changes.")

    add_styled_table(doc, ["Change", "From", "To", "Rationale"],
        [["Weight decay", "1e-4", "0.05 (500x)", "Primary overfitting counter"],
         ["Encoder freezing", "None", "Phase 1/2", "Protect pretrained features"],
         ["Encoder LR", "Same", "1e-5 (30x lower)", "Preserve ImageNet knowledge"],
         ["Effective batch", "2", "8", "Stable rare-class gradients"],
         ["Class weights", "Freq. inverse", "Log-inverse (all)", "Better rare-class weighting"],
         ["Label smoothing", "None", "0.05", "Reduce overconfidence"],
         ["Early stopping", "None", "Patience=10", "Prevent wasted compute"],
         ["Primary metric", "mIoU", "Damage macro F1", "Excludes background"]])

    add_styled_table(doc, ["Metric", "Best Value", "vs. Exp 2"],
        [["Val mIoU", "0.475", "+0.081 (+20.6%)"], ["Val Mean Dice", "0.603", "+0.122 (+25.4%)"],
         ["Val Damage F1", "0.510", "(new metric)"]])

    doc.add_paragraph("Per-Class F1 at Best Epoch (32):")
    add_styled_table(doc, ["Class", "F1 Score"],
        [["No Damage", "0.672"], ["Minor Damage", "0.283"], ["Major Damage", "0.461"], ["Destroyed", "0.623"]])

    p = doc.add_paragraph()
    run = p.add_run("Outcome: "); run.bold = True
    p.add_run(
        "The single largest improvement -- zero architecture changes. 500x weight decay eliminated "
        "overfitting (train-val gap from 3.8:1 to 1:1). Two-phase training stabilised learning.")
    add_figure(doc, "A2_exp2v3_overfitting_fix.png",
        "Figure 5: Exp 2 (3.8x gap) vs Exp 3 (1:1 ratio). The 500x weight decay was the primary driver.", width=6.2)

    p = doc.add_paragraph()
    run = p.add_run("Takeaway: "); run.bold = True
    p.add_run("Training strategy alone: 60% of total improvement. But Minor Damage F1=0.283 -- ResNet34 lacked capacity.")

    # ── Exp 4 ──
    doc.add_heading("4.4 Experiment 4: ResNet50 Backbone -- A Mixed Result", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: "); run.bold = True
    p.add_run("Increase encoder capacity for fine-grained damage distinction, particularly Minor Damage.")

    add_styled_table(doc, ["Change", "From", "To", "Rationale"],
        [["Backbone", "ResNet34 (21M)", "ResNet50 (25M)", "4x wider bottleneck layers"],
         ["Decoder channels", "[256..32]", "[512..64]", "Match wider encoder"],
         ["Batch size", "2", "16", "Utilise P100 VRAM"]])

    add_styled_table(doc, ["Metric", "Best Value", "vs. Exp 3"],
        [["Val mIoU", "0.488", "+0.013 (+2.7%)"], ["Val Damage F1", "0.521", "+0.011 (+2.2%)"]])

    add_styled_table(doc, ["Class", "Exp 3", "Exp 4", "Change"],
        [["No Damage", "0.672", "0.709", "+0.037"], ["Minor Damage", "0.283", "0.433", "+0.150 (+53%)"],
         ["Major Damage", "0.461", "0.384", "-0.077 (-17%)"], ["Destroyed", "0.623", "0.557", "-0.066"]])

    p = doc.add_paragraph()
    run = p.add_run("Outcome: "); run.bold = True
    p.add_run(
        "Minor Damage +53%, but Major Damage and Destroyed regressed. The encoder barely adapted "
        "(encoder_lr=1e-5, only 2 layers unfrozen).")
    add_figure(doc, "A3_exp3v4_perclass.png",
        "Figure 6: Per-class F1. ResNet50 helped Minor Damage (+53%) but Major Damage regressed (-17%).", width=5.5)

    p = doc.add_paragraph()
    run = p.add_run("Takeaway: "); run.bold = True
    p.add_run("A bigger backbone is not automatically better -- you have to let it learn.")

    # ── Exp 5 ──
    doc.add_heading("4.5 Experiment 5: Tuned Phase 2 + Focal Loss", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Goal: "); run.bold = True
    p.add_run("Fix Exp 4's insufficient encoder adaptation. Let ResNet50 actually learn satellite features.")

    add_styled_table(doc, ["Change", "From", "To", "Rationale"],
        [["Encoder LR", "1e-5", "5e-5 (5x)", "Faster encoder adaptation"],
         ["Freeze epochs", "15", "10", "More Phase 2 time"],
         ["Unfrozen layers", "2", "4", "Deeper adaptation"],
         ["Loss", "CE + Dice", "Focal + Dice", "Focus on hard boundaries"],
         ["Epochs", "50", "60", "More convergence time"]])

    add_styled_table(doc, ["Metric", "Best Value", "vs. Exp 4"],
        [["Val mIoU", "0.494", "+0.006"], ["Val Damage F1", "0.533", "+0.012"]])

    add_styled_table(doc, ["Class", "Exp 4", "Exp 5", "Change"],
        [["No Damage", "0.709", "0.683", "-0.026"], ["Minor Damage", "0.433", "0.466", "+0.033"],
         ["Major Damage", "0.384", "0.427", "+0.043"], ["Destroyed", "0.557", "0.556", "-0.001"]])

    p = doc.add_paragraph()
    run = p.add_run("Outcome: "); run.bold = True
    p.add_run(
        "Best results despite only 57% of planned epochs (34/60). Higher encoder LR + Focal Loss "
        "recovered Major Damage F1. Minor Damage F1=0.466 exceeds competition winner range (0.30-0.45).")
    add_figure(doc, "A5_exp5_deepdive.png",
        "Figure 7: Exp 5 deep-dive. Top: Damage F1 by phase. Bottom: Per-class F1 trajectories.", width=6.0)
    add_figure(doc, "A4_exp4v5_phase2_adaptation.png",
        "Figure 8: Exp 4 vs 5 Phase 2 comparison. More aggressive fine-tuning unlocks ResNet50.", width=5.8)

    p = doc.add_paragraph()
    run = p.add_run("Compute constraint: ")
    run.bold = True
    run.font.color.rgb = RGBColor(13, 71, 161)
    p.add_run(
        "Training terminated at epoch 34/60 due to exhausted GPU quota. SUTD AI cluster also down. "
        "Early stopping never triggered -- results likely underestimate full potential.")

    # ── Results Summary ──
    doc.add_heading("4.6 Results Summary", level=2)
    add_styled_table(doc,
        ["", "Exp 1", "Exp 2", "Exp 3", "Exp 4", "Exp 5"],
        [["Key Change", "Baseline", "More epochs", "Train. overhaul", "ResNet50", "Focal+Tuned"],
         ["Backbone", "ResNet34", "ResNet34", "ResNet34", "ResNet50", "ResNet50"],
         ["Epochs", "5", "35", "42", "32", "34/60*"],
         ["Val mIoU", "0.359", "0.394", "0.475", "0.488", "0.494"],
         ["Val Dice", "0.443", "0.481", "0.603", "0.616", "0.625"],
         ["Val Dmg F1", "--", "--", "0.510", "0.521", "0.533"],
         ["Overfitting", "Moderate", "Severe", "Controlled", "Controlled", "Controlled"]])
    p = doc.add_paragraph("*Exp 5 stopped at 34/60 epochs due to GPU quota. Training was still improving.")
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9)

    add_figure(doc, "B5_combined_progression.png",
        "Figure 9: mIoU and Damage F1 progression across all 5 experiments (+37% total mIoU gain).", width=6.0)
    add_figure(doc, "D2_improvement_closeup.png",
        "Figure 16: Close-up showing progressive improvement from Exp 1 to Exp 5.", width=6.2)

    # ── SOTA ──
    doc.add_heading("4.7 Comparison to State-of-the-Art", level=2)
    add_styled_table(doc, ["Model", "Damage F1", "Data", "Ensemble"],
        [["1st place (vdurnov)", "~0.75", "Full xBD (22K)", "12 models"],
         ["2nd place (selimsef)", "~0.72", "Full xBD", "2 models"],
         ["Single-model baseline", "~0.625", "Full xBD", "1 model"],
         ["Ours (Exp 5)", "0.533", "Tier 3 (2.2K)", "1 model"]])
    add_figure(doc, "B4_sota_comparison.png",
        "Figure 10: SOTA comparison. Gap due to 10x less data, no ensembling, incomplete training.", width=5.8)

    # ╔════════════════════════════════════════════════════════════╗
    # ║  5. ANALYSIS                                               ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("5. Analysis", level=1)

    doc.add_heading("5.1 What Made the Biggest Difference", level=2)
    add_figure(doc, "B1_miou_waterfall.png",
        "Figure 11: Cumulative mIoU contribution. Exp 3 (training strategy) accounts for 60% of total gain.", width=6.0)
    add_styled_table(doc, ["Rank", "Change", "Experiment", "mIoU Gain", "% of Total"],
        [["1", "Weight decay + two-phase", "Exp 3", "+0.081", "60%"],
         ["2", "More epochs", "Exp 2", "+0.035", "26%"],
         ["3", "ResNet50 backbone", "Exp 4", "+0.013", "10%"],
         ["4", "Focal Loss + tuned FT", "Exp 5", "+0.006", "4%"]])

    doc.add_heading("5.2 The Rare-Class Challenge", level=2)
    add_figure(doc, "B2_rare_class_challenge.png",
        "Figure 12: Per-class F1 across Exp 3/4/5. Minor Damage improved +65% from 0.28 to 0.47.", width=5.8)

    doc.add_heading("5.3 The Encoder Adaptation Arc", level=2)
    add_figure(doc, "B3_encoder_adaptation_arc.png",
        "Figure 13: Per-class F1 trajectories. Major Damage regressed with conservative fine-tuning, recovered with 5x encoder LR.", width=5.5)
    p = doc.add_paragraph()
    run = p.add_run("Lesson: A bigger backbone with conservative fine-tuning can be worse than a smaller backbone that is fully adapted.")
    run.bold = True

    doc.add_heading("5.4 Loss Curves Across All Experiments", level=2)
    add_figure(doc, "C4_loss_landscape_all.png",
        "Figure 14: Loss curves. Exp 1-2: low loss but overfitting. Exp 3-5: higher loss (class weights) but better generalisation.", width=6.0)

    doc.add_heading("5.5 Failure Cases and Limitations", level=2)
    for item in [
        "Minor vs. Major confusion: boundary between adjacent damage levels is subjective -- human annotators also disagree.",
        "Small buildings: few pixels at 512x512 are hard to classify. Higher resolution would help.",
        "Underrepresented disaster types: fewer training samples = worse generalisation.",
        "Phase 2 validation instability: occasional loss spikes from aggressive encoder fine-tuning.",
        "Incomplete training: Exp 5 completed only 57% of epochs. Results are a lower bound.",
    ]:
        doc.add_paragraph(item, style='List Bullet')
    add_figure(doc, "D3_failure_cases.png",
        "Figure 17: Failure analysis. Error maps show correct (green), wrong class (red), missed buildings (orange).", width=6.0)

    # ╔════════════════════════════════════════════════════════════╗
    # ║  6. REPRODUCIBILITY                                        ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("6. Reproducibility", level=1)

    doc.add_heading("6.1 Environment", level=2)
    doc.add_paragraph(
        "Python 3.11+, PyTorch >= 2.0 (with CUDA), torchvision, opencv-python, scikit-image, "
        "scikit-learn, shapely, matplotlib, numpy, tqdm.")

    doc.add_heading("6.2 Dataset Setup", level=2)
    for item in [
        "Download xView2 Tier 3 from Kaggle",
        "Place in: data/train/train/images/ and data/train/train/labels/",
        "Update data_root in config JSON if using a different path",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("6.3 Training from Scratch", level=2)
    for item in [
        "Open Kaggle-latest.ipynb (Kaggle) or training_script.ipynb (local)",
        "Set CONFIG_PATH to desired experiment config",
        "Run all cells. Checkpoints, metrics, and config saved automatically.",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("6.4 Loading a Trained Model", level=2)
    doc.add_paragraph(
        "model = SiameseResUNet(backbone='resnet50', num_classes=5, decoder_channels=[512, 256, 128, 64])\n"
        "checkpoint = torch.load('training_runs/siamese_resunet_xview2_4/best_model.pt', map_location='cpu')\n"
        "model.load_state_dict(checkpoint['model_state_dict'])\n"
        "model.eval()\n"
        "pred = model(pre_image, post_image).argmax(dim=1)")

    doc.add_heading("6.5 Recreating Figures", level=2)
    doc.add_paragraph("python generate_figures.py        # Figures A1-A5, B1-B5")
    doc.add_paragraph("python generate_extra_figures.py   # Figures C1-C4")
    doc.add_paragraph("python generate_prediction_figures.py  # Figures D1-D3")

    # ╔════════════════════════════════════════════════════════════╗
    # ║  7. CONCLUSION                                             ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "We developed a Siamese ResNet U-Net for building damage segmentation, achieving damage macro "
        "F1 of 0.533 and mIoU of 0.494 through 5 systematic experiments.")

    findings = [
        ("1. Training strategy matters more than architecture.", "Exp 3 delivered +20.6% mIoU with zero architecture changes. Weight decay (0.05) and two-phase training drove 60% of total improvement."),
        ("2. Honest metrics reveal the real challenge.", "Switching from pixel accuracy to damage F1 exposed Minor Damage (F1=0.28) as the bottleneck."),
        ("3. Backbone capacity helps, but only with proper fine-tuning.", "ResNet50 improved Minor Damage +0.150 but regressed Major Damage when fine-tuned conservatively. 5x encoder LR + 4 layers (Exp 5) recovered it."),
        ("4. Focal Loss focuses on what matters.", "Switching to Focal Loss (gamma=2) recovered Major Damage F1 and improved Minor Damage further."),
        ("5. The gap to SOTA is explainable.", "0.533 (single model, 10% data) vs 0.75 (12-model ensemble, full data) -- resource constraints, not architecture."),
    ]
    for title, body in findings:
        p = doc.add_paragraph()
        run = p.add_run(title + " ")
        run.bold = True
        p.add_run(body)

    doc.add_heading("Future Work", level=2)
    for item in [
        "Complete Exp 5 training (26 remaining epochs)",
        "Implement test-time augmentation (+1-3% free gain)",
        "Add auxiliary building localisation head (multi-task learning)",
        "Train on full xBD dataset (~22,000 pairs)",
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ╔════════════════════════════════════════════════════════════╗
    # ║  8. TEAM CONTRIBUTIONS                                     ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("8. Team Contributions", level=1)
    add_styled_table(doc, ["Member", "Contributions"],
        [["[Name 1]", "[Contributions]"], ["[Name 2]", "[Contributions]"],
         ["[Name 3]", "[Contributions]"], ["[Name 4]", "[Contributions]"]])

    # ╔════════════════════════════════════════════════════════════╗
    # ║  REFERENCES                                                ║
    # ╚════════════════════════════════════════════════════════════╝
    doc.add_heading("References", level=1)
    refs = [
        '[1] Gupta, R., et al. "xBD: A Dataset for Assessing Building Damage from Satellite Imagery." arXiv:1911.09296, 2019.',
        '[2] Ronneberger, O., Fischer, P., & Brox, T. "U-Net: Convolutional Networks for Biomedical Image Segmentation." MICCAI 2015.',
        '[3] He, K., et al. "Deep Residual Learning for Image Recognition." CVPR 2016.',
        '[4] Lin, T.-Y., et al. "Focal Loss for Dense Object Detection." ICCV 2017.',
        '[5] xView2 1st Place Solution: github.com/vdurnov/xview2_1st_place_solution',
        '[6] xView2 2nd Place Solution: github.com/selimsef/xview2_solution',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)

    # ── Save ──
    doc.save(str(OUT))
    print(f"Report saved to: {OUT}")


if __name__ == "__main__":
    build()
